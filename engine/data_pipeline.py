"""
BlankWhale v2 — Smart Document Extraction Pipeline
Multi-format intelligent extraction that converts PDFs, DOCX, TXT, CSV,
images, and scans into clean markdown + structured QA-ready data.

Replaces the old raw-chunk approach with structure-aware extraction
that preserves headers, tables, lists, and semantic boundaries.
"""

import json
import csv
import re
import logging
import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any
from .debug_logger import log_function, logger

logger = logging.getLogger("blankwhale.pipeline")


# ============================================================
# Configuration
# ============================================================

@dataclass
class DataConfig:
    """Configuration for data extraction and tokenization."""
    train_file: str = "./data/train.jsonl"
    eval_file: Optional[str] = None
    format: str = "alpaca"            # alpaca | sharegpt | completion
    max_seq_length: int = 2048
    num_workers: int = 4
    shuffle: bool = True
    validation_split: float = 0.1


@dataclass
class ExtractionResult:
    """Result from extracting a single document."""
    source: str                        # Original file path
    format: str                        # File type (pdf, docx, txt, csv, image)
    markdown: str                      # Clean extracted markdown text
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: int = 0
    word_count: int = 0
    has_tables: bool = False
    has_images: bool = False


# ============================================================
# Multi-Format Loader
# ============================================================

@log_function
def extract_document(path: str, max_pages: Optional[int] = None) -> ExtractionResult:
    """
    Intelligently extract text from any supported document format.
    Returns clean markdown with preserved structure.
    
    Supported formats: PDF, DOCX, TXT, CSV, MD, JSON, JSONL, images (OCR)
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    
    ext = p.suffix.lower()
    
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,
        ".txt": _extract_text,
        ".md": _extract_text,
        ".csv": _extract_csv,
        ".json": _extract_json,
        ".jsonl": _extract_jsonl,
        ".png": _extract_image,
        ".jpg": _extract_image,
        ".jpeg": _extract_image,
        ".tiff": _extract_image,
        ".bmp": _extract_image,
    }
    
    extractor = extractors.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file format: {ext}")
    
    logger.info(f"Extracting {ext} document: {p.name}")
    return extractor(p, max_pages=max_pages) if ext == ".pdf" else extractor(p)


def _extract_pdf(path: Path) -> ExtractionResult:
    """
    Extract structured text from PDF.
    
    Strategy:
    1. Use PyMuPDF (fitz) first - it's much faster for text-heavy documents.
    2. Use pdfplumber only for specialized layout/table parsing if fitz is not enough.
    """
    # For general extraction, fitz is the winner on speed
    try:
        return _extract_pdf_pymupdf(path)
    except Exception as e:
        logger.warning(f"PyMuPDF failed ({e}), trying pdfplumber...")
        return _extract_pdf_pdfplumber(path)

def _extract_pdf_pdfplumber(path: Path) -> ExtractionResult:
    """Extract using pdfplumber (fallback)."""
    markdown_parts = []
    page_count = 0
    has_tables = False
    
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            
            for i, page in enumerate(pdf.pages):
                page_md = []
                
                # Extract tables first
                tables = page.extract_tables()
                if tables:
                    has_tables = True
                    for table in tables:
                        page_md.append(_table_to_markdown(table))
                
                # Extract text (excluding table areas for cleaner output)
                text = page.extract_text(
                    x_tolerance=2,
                    y_tolerance=3,
                    layout=True
                )
                
                if text:
                    cleaned = _clean_extracted_text(text)
                    if cleaned:
                        page_md.append(cleaned)
                
                if page_md:
                    markdown_parts.append(f"\n\n".join(page_md))
        
        logger.info(f"pdfplumber extracted {page_count} pages")
        
    except ImportError:
        logger.warning("pdfplumber not found, falling back to PyMuPDF")
        return _extract_pdf_pymupdf(path)
    except Exception as e:
        logger.warning(f"pdfplumber failed ({e}), falling back to PyMuPDF")
        return _extract_pdf_pymupdf(path)
    
    full_markdown = "\n\n---\n\n".join(markdown_parts)
    full_markdown = _normalize_markdown(full_markdown)
    
    return ExtractionResult(
        source=str(path),
        format="pdf",
        markdown=full_markdown,
        pages=page_count,
        word_count=len(full_markdown.split()),
        has_tables=has_tables,
    )


def _extract_pdf_pymupdf(path: Path, max_pages: Optional[int] = None) -> ExtractionResult:
    """Fallback PDF extraction using PyMuPDF (fitz)."""
    try:
        import fitz
    except ImportError:
        raise ImportError("Neither pdfplumber nor PyMuPDF found. Install: pip install pdfplumber pymupdf")
    
    doc = fitz.open(path)
    markdown_parts = []
    
    for i, page in enumerate(doc):
        if max_pages and i >= max_pages:
            break
        blocks = page.get_text("dict")["blocks"]
        page_lines = []
        
        for block in blocks:
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    line_text = " ".join(s["text"].strip() for s in spans if s["text"].strip())
                    
                    if not line_text:
                        continue
                    
                    # Detect headers by font size
                    max_size = max((s.get("size", 12) for s in spans), default=12)
                    is_bold = any(s.get("flags", 0) & 2 for s in spans)
                    
                    if max_size >= 18:
                        page_lines.append(f"# {line_text}")
                    elif max_size >= 14 or is_bold:
                        page_lines.append(f"## {line_text}")
                    else:
                        page_lines.append(line_text)
        
        if page_lines:
            markdown_parts.append("\n".join(page_lines))
    
    doc.close()
    
    full_markdown = "\n\n---\n\n".join(markdown_parts)
    full_markdown = _normalize_markdown(full_markdown)
    
    return ExtractionResult(
        source=str(path),
        format="pdf",
        markdown=full_markdown,
        pages=len(doc) if hasattr(doc, '__len__') else 0,
        word_count=len(full_markdown.split()),
    )


def _extract_docx(path: Path) -> ExtractionResult:
    """Extract text from DOCX with headings, lists, and tables."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not found. Install: pip install python-docx")
    
    doc = Document(path)
    markdown_parts = []
    has_tables = False
    
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        
        if tag == "p":
            # Process paragraphs
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    style = para.style.name.lower() if para.style else ""
                    
                    if "heading 1" in style:
                        markdown_parts.append(f"# {text}")
                    elif "heading 2" in style:
                        markdown_parts.append(f"## {text}")
                    elif "heading 3" in style:
                        markdown_parts.append(f"### {text}")
                    elif "list" in style:
                        markdown_parts.append(f"- {text}")
                    else:
                        markdown_parts.append(text)
                    break
        
        elif tag == "tbl":
            has_tables = True
            for table in doc.tables:
                if table._element is element:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    if rows:
                        markdown_parts.append(_table_to_markdown(rows))
                    break
    
    full_markdown = "\n\n".join(markdown_parts)
    full_markdown = _normalize_markdown(full_markdown)
    
    return ExtractionResult(
        source=str(path),
        format="docx",
        markdown=full_markdown,
        word_count=len(full_markdown.split()),
        has_tables=has_tables,
    )


def _extract_text(path: Path) -> ExtractionResult:
    """Extract plain text or markdown files."""
    text = path.read_text(encoding="utf-8", errors="replace")
    text = _normalize_markdown(text)
    
    return ExtractionResult(
        source=str(path),
        format="txt",
        markdown=text,
        word_count=len(text.split()),
    )


def _extract_csv(path: Path) -> ExtractionResult:
    """Convert CSV to markdown table."""
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(row)
    
    if not rows:
        return ExtractionResult(source=str(path), format="csv", markdown="")
    
    md = _table_to_markdown(rows)
    
    return ExtractionResult(
        source=str(path),
        format="csv",
        markdown=md,
        word_count=len(md.split()),
        has_tables=True,
    )


def _extract_json(path: Path) -> ExtractionResult:
    """Extract JSON into readable markdown."""
    data = json.loads(path.read_text(encoding="utf-8"))
    
    if isinstance(data, list):
        md_parts = []
        for i, item in enumerate(data[:100]):  # Cap at 100 items
            md_parts.append(f"### Record {i+1}\n```json\n{json.dumps(item, indent=2, ensure_ascii=False)}\n```")
        md = "\n\n".join(md_parts)
    else:
        md = f"```json\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```"
    
    return ExtractionResult(
        source=str(path),
        format="json",
        markdown=md,
        word_count=len(md.split()),
    )


def _extract_jsonl(path: Path) -> ExtractionResult:
    """Extract JSONL into readable markdown."""
    lines = path.read_text(encoding="utf-8").strip().split("\n")
    md_parts = []
    
    for i, line in enumerate(lines[:100]):  # Cap at 100
        try:
            item = json.loads(line)
            md_parts.append(f"### Record {i+1}\n```json\n{json.dumps(item, indent=2, ensure_ascii=False)}\n```")
        except json.JSONDecodeError:
            continue
    
    md = "\n\n".join(md_parts)
    
    return ExtractionResult(
        source=str(path),
        format="jsonl",
        markdown=md,
        word_count=len(md.split()),
    )


def _extract_image(path: Path) -> ExtractionResult:
    """Extract text from images using OCR (pytesseract)."""
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        raise ImportError("OCR deps missing. Install: pip install Pillow pytesseract")
    
    image = Image.open(path)
    text = pytesseract.image_to_string(image)
    text = _normalize_markdown(text.strip())
    
    return ExtractionResult(
        source=str(path),
        format="image",
        markdown=text,
        word_count=len(text.split()),
        has_images=True,
    )


# ============================================================
# Semantic Chunking
# ============================================================

def chunk_by_semantic_boundaries(
    markdown: str,
    chunk_size: int = 1024,
    overlap: int = 128,
) -> List[str]:
    """
    Split markdown into chunks using semantic boundaries.
    
    Strategy:
    1. Split on markdown headers (# ## ###)
    2. Then split on paragraph boundaries (double newlines)
    3. Only fall back to character-level splits for very long paragraphs
    
    This preserves context much better than raw character chunking.
    """
    if not markdown or len(markdown) <= chunk_size:
        return [markdown] if markdown else []
    
    # Split on headers first
    sections = re.split(r'\n(?=#{1,3}\s)', markdown)
    
    chunks = []
    current_chunk = ""
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        # If the section fits, add it
        if len(current_chunk) + len(section) + 2 <= chunk_size:
            current_chunk = (current_chunk + "\n\n" + section).strip()
        else:
            # Save current chunk
            if current_chunk:
                chunks.append(current_chunk)
            
            # If section itself is too long, split by paragraphs
            if len(section) > chunk_size:
                sub_chunks = _split_long_section(section, chunk_size, overlap)
                chunks.extend(sub_chunks[:-1])
                current_chunk = sub_chunks[-1] if sub_chunks else ""
            else:
                # Start new chunk with overlap from previous
                if chunks:
                    overlap_text = chunks[-1][-overlap:]
                    current_chunk = overlap_text + "\n\n" + section
                else:
                    current_chunk = section
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks


def _split_long_section(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Split a long section by paragraphs, then by sentences."""
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        if len(current) + len(para) + 2 <= chunk_size:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            
            if len(para) > chunk_size:
                # Last resort: split by sentences
                sentences = re.split(r'(?<=[.!?])\s+', para)
                current = ""
                for sent in sentences:
                    if len(current) + len(sent) + 1 <= chunk_size:
                        current = (current + " " + sent).strip()
                    else:
                        if current:
                            chunks.append(current)
                        current = sent
            else:
                current = para
    
    if current.strip():
        chunks.append(current.strip())
    
    return chunks


# ============================================================
# Helpers
# ============================================================

def _table_to_markdown(rows: List[List[str]]) -> str:
    """Convert a list of rows into a markdown table."""
    if not rows:
        return ""
    
    # Clean cells
    clean_rows = []
    for row in rows:
        clean_rows.append([str(cell).strip().replace("|", "\\|") if cell else "" for cell in row])
    
    # Header
    header = "| " + " | ".join(clean_rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * len(clean_rows[0])) + " |"
    
    # Body
    body_lines = []
    for row in clean_rows[1:]:
        # Pad row to match header length
        while len(row) < len(clean_rows[0]):
            row.append("")
        body_lines.append("| " + " | ".join(row[:len(clean_rows[0])]) + " |")
    
    return "\n".join([header, separator] + body_lines)


def _clean_extracted_text(text: str) -> str:
    """Clean raw extracted text: fix encoding, remove artifacts."""
    if not text:
        return ""
    
    # Remove page numbers / headers that are just numbers
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    # Remove excessive whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    # Normalize line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def _normalize_markdown(text: str) -> str:
    """Final normalization pass on extracted markdown."""
    if not text:
        return ""
    
    # Fix broken unicode
    text = text.encode("utf-8", errors="replace").decode("utf-8")
    # Remove null bytes
    text = text.replace("\x00", "")
    # Normalize whitespace
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    text = re.sub(r' {3,}', '  ', text)
    
    return text.strip()


# ============================================================
# Legacy Compatibility — load_raw_data / format functions
# ============================================================

def load_raw_data(file_path: str, chunk_size: int = 1024, overlap: int = 128) -> List[str]:
    """
    High-level function: extract a document and return semantic chunks.
    This replaces the old raw chunking approach entirely.
    """
    result = extract_document(file_path)
    chunks = chunk_by_semantic_boundaries(result.markdown, chunk_size, overlap)
    logger.info(f"Extracted {len(chunks)} chunks from {Path(file_path).name} "
                f"({result.word_count} words, {result.pages} pages)")
    return chunks


def format_for_training(
    chunks: List[str],
    format_type: str = "completion",
    system_prompt: str = "You are a helpful assistant.",
) -> List[Dict[str, str]]:
    """
    Convert text chunks into training-ready format.
    
    Formats:
    - completion: {"text": "..."}
    - alpaca: {"instruction": "...", "input": "", "output": "..."}
    - sharegpt: {"conversations": [{"from": "human", ...}, {"from": "gpt", ...}]}
    """
    records = []
    
    for chunk in chunks:
        if not chunk.strip():
            continue
        
        if format_type == "completion":
            records.append({"text": chunk})
        
        elif format_type == "alpaca":
            records.append({
                "instruction": "Read and understand the following content:",
                "input": "",
                "output": chunk,
            })
        
        elif format_type == "sharegpt":
            records.append({
                "conversations": [
                    {"from": "human", "value": "Explain the following content:"},
                    {"from": "gpt", "value": chunk},
                ]
            })
    
    return records


def save_training_data(records: List[Dict], output_path: str = "./data/train.jsonl"):
    """Save formatted records to JSONL."""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, "w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    
    logger.info(f"Saved {len(records)} training records to {output_path}")

@log_function
def preprocess_data(config: DataConfig):
    """Load prepared dataset for training, with synthetic fallback."""
    from datasets import load_dataset
    import os
    
    # Check if file exists and has content
    file_exists = os.path.exists(config.train_file)
    is_empty = file_exists and os.path.getsize(config.train_file) < 10

    if not file_exists or is_empty:
        logger.warning(f"Training file {config.train_file} not found or empty. Creating synthetic dataset.")
        create_synthetic_dataset(config.train_file, format_type=config.format)

    data_files = {"train": config.train_file}
    if config.eval_file and os.path.exists(config.eval_file):
        data_files["eval"] = config.eval_file
            
    dataset = load_dataset("json", data_files=data_files)
    
    if config.shuffle and "train" in dataset:
        dataset["train"] = dataset["train"].shuffle(seed=42)
        
    return dataset

@log_function
def create_synthetic_dataset(output_path: str, format_type: str = "alpaca"):
    """Create a minimal 10-example synthetic dataset for debugging."""
    logger.info(f"Creating synthetic dataset at {output_path} (format: {format_type})")
    examples = []
    
    topics = [
        "What is BlankWhale?", "How to train a model", "GPU detection", 
        "Data privacy", "Local AI benefits", "Robotics simulation",
        "Fine-tuning techniques", "LoRA vs QLoRA", "Markdown extraction", 
        "Model quantization"
    ]
    
    for i, topic in enumerate(topics):
        if format_type == "alpaca":
            examples.append({
                "instruction": f"Explain {topic} in the context of BlankWhale.",
                "input": "",
                "output": f"BlankWhale provides expert tools for {topic}, ensuring privacy and local execution."
            })
        else:
            examples.append({
                "text": f"### Topic: {topic}\nBlankWhale is an expert local AI studio that handles {topic} efficiently on consumer hardware."
            })
            
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        for ex in examples:
            f.write(json.dumps(ex, ensure_ascii=False) + "\n")
    
    logger.info(f"Successfully created 10 synthetic examples.")
